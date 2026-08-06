import io
import os
import re
import tempfile
import zipfile
import subprocess
import shutil
from typing import List, Tuple

# Document Parsing Dependencies
from docx import Document
import pdfplumber
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from lxml import etree

LIBREOFFICE_BIN = shutil.which("soffice") or shutil.which("libreoffice")
LIBREOFFICE_OK = LIBREOFFICE_BIN is not None

# Formats LibreOffice converts to PDF directly
LO_DIRECT_EXTS = {
    "doc", "docx", "dot", "dotx", "odt", "ott", "rtf", "txt", "md",
    "html", "htm", "xhtml", "wps",
}
# Formats requiring custom reading
CUSTOM_ONLY_EXTS = {"idml", "xml", "dita", "epub"}
UNSUPPORTED_EXTS = {"indd", "pages"}


# =========================================================================
# READERS — every format converges to List[Tuple[role, text]]
# =========================================================================

def read_docx(file_bytes: bytes) -> List[Tuple[str, str]]:
    doc = Document(io.BytesIO(file_bytes))
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        role = "heading" if ("heading" in style or "title" in style) else "body"
        blocks.append((role, text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    blocks.append(("body", text))
    return blocks

def read_pdf(file_bytes: bytes) -> List[Tuple[str, str]]:
    blocks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text(layout=False) or ""
            # Clean up single newlines that are just line breaks within a paragraph
            # But keep double newlines which indicate true paragraph separation
            text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
            for para in re.split(r'\n+', text):
                para = para.strip()
                if para:
                    blocks.append(("body", para))
    return blocks

def read_epub(file_bytes: bytes) -> List[Tuple[str, str]]:
    with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as f:
        f.write(file_bytes)
        path = f.name
    try:
        book = epub.read_epub(path)
        blocks = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            for tag in soup.find_all(("p", "h1", "h2", "h3", "h4", "li")):
                text = tag.get_text(separator=" ", strip=True)
                if text:
                    role = "heading" if tag.name.startswith("h") else "body"
                    blocks.append((role, text))
        return blocks
    finally:
        os.unlink(path)

def read_idml(file_bytes: bytes) -> List[Tuple[str, str]]:
    blocks = []
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        story_files = [n for n in z.namelist() if n.startswith("Stories/Story_") and n.endswith(".xml")]
        for name in sorted(story_files):
            tree = etree.fromstring(z.read(name))
            for para_range in tree.iter():
                if para_range.tag.endswith("ParagraphStyleRange"):
                    parts = [t.text for t in para_range.iter() if t.tag.endswith("Content") and t.text]
                    text = "".join(parts).strip()
                    if text:
                        blocks.append(("body", text))
    return blocks

def read_txt(file_bytes: bytes) -> List[Tuple[str, str]]:
    text = file_bytes.decode("utf-8", errors="replace")
    blocks = []
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if para:
            role = "heading" if len(para) < 60 and para.isupper() else "body"
            blocks.append((role, para))
    return blocks

def read_md(file_bytes: bytes) -> List[Tuple[str, str]]:
    text = file_bytes.decode("utf-8", errors="replace")
    blocks = []
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if not para:
            continue
        if para.startswith("#"):
            blocks.append(("heading", para.lstrip("#").strip()))
        else:
            blocks.append(("body", re.sub(r"[*_`]", "", para)))
    return blocks

def read_html(file_bytes: bytes) -> List[Tuple[str, str]]:
    soup = BeautifulSoup(file_bytes, "html.parser")
    blocks = []
    for tag in soup.find_all(("p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "td")):
        text = tag.get_text(separator=" ", strip=True)
        if text:
            role = "heading" if tag.name.startswith("h") else "body"
            blocks.append((role, text))
    return blocks

def read_xml_dita(file_bytes: bytes) -> List[Tuple[str, str]]:
    tree = etree.fromstring(file_bytes)
    blocks = []
    for elem in tree.iter():
        text = (elem.text or "").strip()
        if not text:
            continue
        tag = etree.QName(elem).localname.lower()
        role = "heading" if tag in ("title", "head", "shorttitle") else "body"
        blocks.append((role, text))
    return blocks

READERS = {
    "docx": read_docx,
    "pdf": read_pdf,
    "epub": read_epub,
    "idml": read_idml,
    "txt": read_txt,
    "md": read_md,
    "html": read_html,
    "htm": read_html,
    "xhtml": read_html,
    "xml": read_xml_dita,
    "dita": read_xml_dita,
}

# =========================================================================
# WRITERS
# =========================================================================

def write_pdf(blocks: List[Tuple[str, str]], title="Converted Document") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER)
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle("HeadingStyle", parent=styles["Heading1"], spaceAfter=10)
    body_style = ParagraphStyle("BodyStyle", parent=styles["Normal"], spaceAfter=8, leading=14)

    story = []
    for role, text in blocks:
        safe_text = (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
        if role == "heading":
            story.append(Paragraph(safe_text, heading_style))
        else:
            story.append(Paragraph(safe_text, body_style))
        story.append(Spacer(1, 14))
    doc.build(story)
    return buf.getvalue()

# =========================================================================
# CORE PIPELINE
# =========================================================================

def convert_via_libreoffice(file_bytes: bytes, ext: str) -> bytes:
    if not LIBREOFFICE_OK:
        raise RuntimeError("LibreOffice not found on PATH")
    with tempfile.TemporaryDirectory() as tmp:
        src_path = os.path.join(tmp, f"input.{ext}")
        with open(src_path, "wb") as f:
            f.write(file_bytes)
        result = subprocess.run(
            [LIBREOFFICE_BIN, "--headless", "--norestore", "--convert-to", "pdf",
             "--outdir", tmp, src_path],
            capture_output=True, text=True, timeout=120,
        )
        out_path = os.path.join(tmp, "input.pdf")
        if result.returncode != 0 or not os.path.exists(out_path):
            raise RuntimeError(f"LibreOffice conversion failed: {result.stdout} {result.stderr}")
        with open(out_path, "rb") as f:
            return f.read()

def normalize_to_pdf(file_bytes: bytes, ext: str) -> bytes:
    """Converts any supported input format to a normalized PDF byte stream."""
    ext = ext.lower()
    if ext == "pdf":
        return file_bytes

    if ext in UNSUPPORTED_EXTS:
        raise RuntimeError(f".{ext} is not supported directly. Please convert to a supported format first.")

    if ext in LO_DIRECT_EXTS and LIBREOFFICE_OK:
        try:
            return convert_via_libreoffice(file_bytes, ext)
        except Exception:
            pass # Fall through to block reader

    if ext in READERS:
        blocks = READERS[ext](file_bytes)
        if not blocks:
            raise RuntimeError("No extractable text found.")
        return write_pdf(blocks)

    raise RuntimeError(f"Format .{ext} not supported.")

def parse_document_blocks(file_bytes: bytes, ext: str) -> List[Tuple[str, str]]:
    """Extracts raw text blocks (role, text) from a document."""
    ext = ext.lower()
    if ext in READERS:
        return READERS[ext](file_bytes)
    
    # If not directly readable by a block reader but LibreOffice can convert it to PDF:
    if ext in LO_DIRECT_EXTS and LIBREOFFICE_OK:
        pdf_bytes = convert_via_libreoffice(file_bytes, ext)
        return read_pdf(pdf_bytes)
        
    raise RuntimeError(f"Cannot extract text blocks directly from .{ext}")

def pdf_to_strings(pdf_bytes: bytes) -> List[str]:
    """Helper to just get plain text lines from a PDF."""
    blocks = read_pdf(pdf_bytes)
    return [text for role, text in blocks]
